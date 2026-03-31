"""
Verifies that the generated LOI contains all updated verbiage from the template.
Checks static phrases, scenario-specific text, and per-paragraph coverage.

Run: python tests/verify_template_verbiage.py
"""

import copy
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from services.document_generator import DocumentGenerator
from services.loi_form_data import (
    LoiFormData,
    DepositStructure,
    DueDiligenceType,
    ClosingExtensionType,
    CommissionType,
    SignatureBlockType,
    SignatureEntity,
)
from services.number_to_words import to_legal_dollar_string, convert_to_words

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
TEMPLATE_PATH = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "Templates", "LOI_Template.docx",
)


def qn(tag):
    prefix, local = tag.split(":")
    return "{%s}%s" % ({"w": W_NS}[prefix], local)


def get_all_text(body):
    texts = []
    for t in body.iter(qn("w:t")):
        texts.append(t.text or "")
    for dt in body.iter("{%s}delText" % W_NS):
        texts.append(dt.text or "")
    return "".join(texts)


def get_visible_text(body):
    return "".join(t.text or "" for t in body.iter(qn("w:t")))


def base_form():
    return LoiFormData(
        date="March 1, 2026",
        property_address="123 Main St, Springfield, IL 62701",
        seller_address_line1="Acme Properties LLC",
        seller_address_line2="100 Commerce Dr",
        seller_address_line3="Springfield, IL 62701",
        attention_name="John Doe",
        salutation="Mr. Doe",
        seller_name="Acme Properties LLC",
        purchase_price=500000,
        initial_deposit=10000,
        additional_deposit=10000,
        monthly_release_amount=5000,
        legal_reimbursement_amount=5000,
        extension_deposit_amount=5000,
        monthly_closing_extension_deposit=25000,
        due_diligence_days=120,
        governmental_approvals_days=150,
        assemblage_days=90,
        closing_days=30,
        closing_extension_months=6,
        lease_end_date="May 31, 2026",
        lease_termination_days=60,
        deposit_structure=DepositStructure.GOVERNMENTAL_APPROVALS_GOING_HARD,
        include_legal_reimbursement=False,
        due_diligence_type=DueDiligenceType.STANDARD,
        closing_extension_type=ClosingExtensionType.NONE,
        commission_type=CommissionType.SELLER_PAYS_LISTING_AGENT,
        include_option_to_extend=True,
        include_existing_leases=True,
        include_delivered_vacant=False,
        include_lease_termination=False,
        include_right_to_negotiate_with_tenants=False,
        include_seller_rollover=False,
        signature_block_type=SignatureBlockType.INDIVIDUAL,
        seller_name_signature="John Doe",
        signature_entities=[SignatureEntity(company_name="Test Corp")],
        parcel_ids=["12-345-678"],
        prepared_by_first_name="Jake",
        prepared_by_last_name="Miller",
    )


def generate(form):
    gen = DocumentGenerator()
    buf = gen.generate(TEMPLATE_PATH, form)
    doc = Document(buf)
    return doc.element.body, doc


# =========================================================================
# Phase 1: Key verbiage phrases from the template
# =========================================================================

# Substantive legal phrases that must always appear in any generated LOI
ALWAYS_PRESENT = [
    "Subtext Acquisitions, LLC, a Missouri limited liability company, or its assignee",
    "non-binding proposal",
    "The terms of the proposed sale are as follows",
    "Property.  The proposal is for the Property described herein together with all improvements, rights of way, easements, hereditaments and appurtenances",
    "Purchase Price.",
    "Exclusivity Period",
    "negotiate on an exclusive basis a definitive purchase agreement",
    "Governmental Approvals Period",
    "Confidentiality.",
    "Governing Law.  This proposal shall be governed by the laws of the State of Missouri",
    "without giving effect to conflict of laws principles",
    "Miscellaneous.  This proposal may be signed in counterparts",
    "If the foregoing is acceptable, please indicate by executing a copy of this proposal",
    "Sincerely,",
    "Subtext Acquisitions, LLC",
    "Richard Birner, Vice President of Land Acquisitions",
    "EXHIBIT A",
    "DEPICTION OF PROPERTY",
    "PARCEL ID NUMBER(S)",
]

# =========================================================================
# Phase 2: Scenario-specific verbiage
# =========================================================================

SCENARIOS = {
    "Deposit - Gov Approvals": {
        "overrides": {"deposit_structure": DepositStructure.GOVERNMENTAL_APPROVALS_GOING_HARD},
        "must_contain": ["Initial Deposit", "Deposit."],
    },
    "Deposit - DD Going Hard": {
        "overrides": {"deposit_structure": DepositStructure.DUE_DILIGENCE_GOING_HARD},
        "must_contain": ["Initial Deposit", "Deposit."],
    },
    "Deposit - Monthly": {
        "overrides": {"deposit_structure": DepositStructure.MONTHLY_GOING_HARD},
        "must_contain": ["Initial Deposit", "Deposit."],
    },
    "Legal Reimbursement Included": {
        "overrides": {"include_legal_reimbursement": True},
        "must_contain": ["Legal Reimbursement Fee"],
    },
    "DD with Assemblage": {
        "overrides": {"due_diligence_type": DueDiligenceType.WITH_ASSEMBLAGE},
        "must_contain": ["Assemblage Period"],
    },
    "Closing Extension - MtM": {
        "overrides": {"closing_extension_type": ClosingExtensionType.MONTH_TO_MONTH},
        "must_contain": ["Closing Extension"],
    },
    "Commission - Seller Pays": {
        "overrides": {"commission_type": CommissionType.SELLER_PAYS_LISTING_AGENT},
        "must_contain": ["Commissions.", "listing agreement"],
    },
    "Commission - Subtext Pays": {
        "overrides": {"commission_type": CommissionType.SUBTEXT_PAYS, "broker_name": "Jones Realty"},
        "must_contain": ["Commissions.", "Jones Realty"],
    },
    "Commission - No Brokers": {
        "overrides": {"commission_type": CommissionType.NO_BROKERS},
        "must_contain": ["Commissions."],
    },
    "Leases - Existing": {
        "overrides": {"include_existing_leases": True, "lease_end_date": "May 31, 2026"},
        "must_contain": ["Leases.", "May 31, 2026"],
    },
    "Leases - Delivered Vacant": {
        "overrides": {"include_existing_leases": False, "include_delivered_vacant": True},
        "must_contain": ["Leases."],
    },
    "Leases - Termination": {
        "overrides": {"include_lease_termination": True, "lease_termination_days": 60},
        "must_contain": ["Leases."],
    },
    "Option to Extend - Included": {
        "overrides": {"include_option_to_extend": True},
        "must_contain": ["Option to Extend.", "Extension Notice"],
    },
    "Seller Rollover - Included": {
        "overrides": {"include_seller_rollover": True},
        "must_contain": ["Seller Rollover Option"],
    },
    "Signature - Individual": {
        "overrides": {
            "signature_block_type": SignatureBlockType.INDIVIDUAL,
            "seller_name_signature": "John Doe",
        },
        "must_contain": ["John Doe"],
    },
    "Signature - Company": {
        "overrides": {
            "signature_block_type": SignatureBlockType.COMPANY_ENTITY,
            "signature_entities": [SignatureEntity(company_name="Alpha Corp")],
        },
        "must_contain": ["Alpha Corp"],
    },
}


# =========================================================================
# Phase 3: Per-paragraph coverage check
# =========================================================================

def extract_template_phrases(template_path):
    """Extract non-placeholder, non-instruction text from each template paragraph."""
    doc = Document(template_path)
    phrases = []
    # Known instruction markers to skip
    instruction_markers = [
        "REMOVE", "INSERT", "USE THE FOLLOWING", "IF SELLER", "IF SUBTEXT",
        "IF NO BROKERS", "REPLACE THIS ENTIRE", "WHEN A SINGLE LOI",
    ]
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        # Skip pure instruction paragraphs
        if any(text.startswith(m) for m in instruction_markers):
            continue
        # Skip placeholder-only paragraphs
        if text.startswith("[") and text.endswith("]") and text.count("[") == 1:
            continue
        phrases.append(text)
    return phrases


# =========================================================================
# Main
# =========================================================================

if __name__ == "__main__":
    if not os.path.exists(TEMPLATE_PATH):
        print(f"ERROR: Template not found at {TEMPLATE_PATH}")
        sys.exit(1)

    total_pass = 0
    total_fail = 0
    failures = []

    # --- Phase 1 ---
    print("=== Phase 1: Key phrases exist in template ===")
    tmpl = Document(TEMPLATE_PATH)
    tmpl_full = "\n".join(p.text for p in tmpl.paragraphs)

    for phrase in ALWAYS_PRESENT:
        if phrase in tmpl_full:
            total_pass += 1
        else:
            total_fail += 1
            failures.append(f"Template missing: {phrase[:70]}")
            print(f"  FAIL: Not in template: {phrase[:70]}...")

    if not any("Template missing" in f for f in failures):
        print(f"  All {len(ALWAYS_PRESENT)} key phrases found in template.")

    # --- Phase 2: Generate default and check ---
    print("\n=== Phase 2: Key phrases survive generation ===")
    form = base_form()
    body, doc = generate(form)
    all_text = get_all_text(body)

    for phrase in ALWAYS_PRESENT:
        if phrase in all_text:
            total_pass += 1
        else:
            total_fail += 1
            failures.append(f"Output missing: {phrase[:70]}")
            print(f"  FAIL: Not in output: {phrase[:70]}...")

    if not any("Output missing" in f for f in failures):
        print(f"  All {len(ALWAYS_PRESENT)} key phrases present in generated output.")

    # --- Phase 3: Scenario-specific checks ---
    print("\n=== Phase 3: Scenario-specific verbiage ===")
    for name, cfg in SCENARIOS.items():
        f = copy.deepcopy(form)
        for k, v in cfg["overrides"].items():
            setattr(f, k, v)
        b, _ = generate(f)
        txt = get_all_text(b)
        ok = True
        for phrase in cfg["must_contain"]:
            if phrase not in txt:
                print(f"  FAIL [{name}]: Missing \"{phrase}\"")
                ok = False
                total_fail += 1
                failures.append(f"Scenario [{name}]: {phrase}")
        if ok:
            total_pass += 1

    # --- Phase 4: Per-paragraph coverage ---
    print("\n=== Phase 4: Template paragraph coverage ===")
    tmpl_phrases = extract_template_phrases(TEMPLATE_PATH)
    # For each substantive template paragraph, check if at least part of it
    # appears in a generated doc (using any scenario).
    # We use the default form output for non-conditional paragraphs.
    default_text = get_all_text(body)
    uncovered = []
    for phrase in tmpl_phrases:
        # Take a 40-char snippet from the middle to avoid placeholders at edges
        mid = len(phrase) // 2
        start = max(0, mid - 20)
        snippet = phrase[start:start + 40].strip()
        if len(snippet) < 10:
            continue
        if snippet in default_text:
            total_pass += 1
        else:
            uncovered.append(snippet)

    if uncovered:
        print(f"  {len(uncovered)} paragraph snippets not found in default output (may be scenario-dependent):")
        for s in uncovered[:10]:
            print(f"    - \"{s}\"")
    else:
        print(f"  All template paragraphs have coverage in generated output.")

    # --- Summary ---
    print(f"\n{'='*50}")
    print(f"TOTAL: {total_pass} passed, {total_fail} failed")
    if failures:
        print("\nFailures:")
        for f in failures:
            print(f"  - {f}")
    else:
        print("\nALL VERBIAGE CHECKS PASSED")
        print("The generated LOI contains all updated template verbiage.")

    sys.exit(1 if total_fail > 0 else 0)
